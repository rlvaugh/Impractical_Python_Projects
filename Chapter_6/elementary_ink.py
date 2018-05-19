""" Hide a message in a docx document using a white font."""
import docx
from docx.shared import RGBColor, Pt

# get text from a fake message & make each line a list item
fake_text = docx.Document('fakeMessage.docx')
fake_list = []
for paragraph in fake_text.paragraphs:
    fake_list.append(paragraph.text)

# get text from real message & make each line a list item
real_text = docx.Document('realMessage_Vig.docx')
real_list = []
for paragraph in real_text.paragraphs:
    if len(paragraph.text) != 0:  # remove blank lines
        real_list.append(paragraph.text)

# load template that sets style, font, margins, etc.
doc = docx.Document('template.docx') 

# add letterhead:
doc.add_heading('Morland Holmes', 0)
subtitle = doc.add_heading('Global Consultanting & Negotiations', 1)
subtitle.alignment = 1 
doc.add_heading('', 1)
doc.add_paragraph('December 17, 2015')
doc.add_paragraph('')

def set_spacing(paragraph):
    """Use docx to set line spacing between paragraphs."""
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)

length_real = len(real_list)
count_real = 0  # index of current line in real (hidden) message

# interleave real and fake message lines
for line in fake_list:
    if count_real < length_real and line == "":
        paragraph = doc.add_paragraph(real_list[count_real])
        paragraph_index = len(doc.paragraphs) - 1
        
        # set real message color to white
        run = doc.paragraphs[paragraph_index].runs[0]   
        font = run.font
        font.color.rgb = RGBColor(255, 255, 255)  # make it red to test
        count_real += 1
        
    else:
        paragraph = doc.add_paragraph(line)
        
    set_spacing(paragraph)

doc.save('ciphertext_message_letterhead.docx')

print("Done")
